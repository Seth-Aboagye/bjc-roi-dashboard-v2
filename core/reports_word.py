import io
import pandas as pd
from docx import Document
from core.metrics import compute_rollups

def build_word_report(d: pd.DataFrame, c: pd.DataFrame, payload: dict, notes: str = "") -> bytes:
    doc = Document()
    doc.add_heading(payload["title"], level=1)
    doc.add_paragraph(f"Generated: {payload['generated_at']}")
    doc.add_paragraph(f"Date range: {payload['filters']['start']} to {payload['filters']['end']}")
    doc.add_paragraph(f"Channels: {', '.join(payload['filters']['channels'])}")
    doc.add_paragraph(f"Campaigns: {', '.join(payload['filters']['campaigns'])}")
    doc.add_paragraph(f"Donor segments: {', '.join(payload['filters']['segments'])}")

    doc.add_heading("Executive KPIs", level=2)
    k = payload["kpis"]
    doc.add_paragraph(f"Total Raised: ${k['total_raised']:,.0f}")
    doc.add_paragraph(f"Total Costs: ${k['total_costs']:,.0f}")
    doc.add_paragraph(f"Net Raised: ${k['net_raised']:,.0f}")
    doc.add_paragraph(f"ROI: {k['roi']*100:,.1f}%")
    doc.add_paragraph(f"Cost to Raise $1: ${k['cost_to_raise_1']:,.2f}")
    doc.add_paragraph(f"Donors: {k['donors']:,} | Gifts: {k['gifts']:,} | Avg Gift: ${k['avg_gift']:,.2f}")

    if notes.strip():
        doc.add_heading("Interpretation / Notes", level=2)
        doc.add_paragraph(notes)

    doc.add_heading("Top Campaigns (by Raised)", level=2)
    top = compute_rollups(d, c, by="campaign_code").sort_values("raised", ascending=False).head(15)

    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    hdr[0].text = "Campaign"
    hdr[1].text = "Raised"
    hdr[2].text = "Costs"
    hdr[3].text = "Net"
    hdr[4].text = "ROI"
    hdr[5].text = "Cost/$"

    for _, r in top.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["campaign_code"])
        row[1].text = f"${r['raised']:,.0f}"
        row[2].text = f"${r['costs']:,.0f}"
        row[3].text = f"${r['net']:,.0f}"
        row[4].text = f"{r['roi']*100:,.1f}%"
        row[5].text = f"${r['cost_to_raise_1']:,.2f}"

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
